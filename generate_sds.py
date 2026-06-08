#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
严格按照SDS模板布局生成TTMS软件设计说明书。
仅补充内容，保留原始模板的所有结构和格式。
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# 全局默认样式
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============================================================
# 辅助函数
# ============================================================
def add_heading_styled(text, level=1, font_name='黑体'):
    """添加标题，使用黑体"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_body_para(text, bold=False, first_line_indent=True, font_size=12, font_name='宋体', alignment=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_plain_para(text, font_size=12, font_name='宋体', bold=False, alignment=None):
    """添加无缩进段落"""
    return add_body_para(text, bold=bold, first_line_indent=False, font_size=font_size, font_name=font_name, alignment=alignment)

def add_blank_para(count=1):
    """添加空行"""
    for _ in range(count):
        doc.add_paragraph()

def set_cell_font(cell, text, font_name='宋体', font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文字格式"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text) if text is not None else ' ')
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    return cell

def add_table(headers, rows, col_widths=None, header_bg='D9E2F3'):
    """添加标准表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, header, bold=True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_bg}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_font(cell, val if val is not None else ' ')
    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row_obj in table.rows:
                row_obj.cells[i].width = Cm(width)
    add_blank_para()
    return table

# ============================================================
# 封面 (第1页)
# ============================================================
add_blank_para(4)

# 项目编号/文档编号/密级
info_table = doc.add_table(rows=3, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (k, v) in enumerate([('项目编号', '            '), ('文档编号', '            '), ('密    级', '            ')]):
    set_cell_font(info_table.rows[i].cells[0], k, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(info_table.rows[i].cells[1], v, alignment=WD_ALIGN_PARAGRAPH.LEFT)
for row in info_table.rows:
    row.cells[0].width = Cm(2.5)

add_blank_para(3)

# [pic] 占位 - 项目Logo位置
add_body_para('[pic]', font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)

add_blank_para(2)

# 项目名称
add_body_para('项目名称', font_size=26, bold=True, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
add_body_para('软件设计说明', font_size=22, bold=True, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)

add_blank_para()
add_body_para('版本：V1.0', font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)

add_blank_para(3)

# 团队成员
add_body_para('团队成员：', font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_table(
    ['学号', '姓名', '角色', '分工'],
    [[' ', ' ', ' ', ' '] for _ in range(6)],
    [3, 3, 3, 7]
)

add_blank_para(2)

add_body_para('软件XXXX班XXXXXX软件公司', font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
add_body_para('二○××年×月', font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)

add_blank_para(4)

# 签署栏
for label in ['拟  制：', '审  核：', '标准化：', '会  签：', '批  准：']:
    add_body_para(f'\n\n{label}\n\n\n', first_line_indent=False)

doc.add_page_break()

# ============================================================
# 文档修改记录 (第2页)
# ============================================================
add_heading_styled('文档修改记录', level=1)
add_table(
    ['版本号', '修改内容描述', '修改人', '日期', '备注'],
    [[' ', ' ', ' ', ' ', ' '] for _ in range(20)],
    [2.5, 6, 2.5, 2.5, 2.5]
)

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
add_heading_styled('目  录', level=1)

toc_items = [
    ('1. 引言', '1'),
    ('    1.1 编写目的', '1'), ('    1.2 项目概述', '1'), ('    1.3 术语定义', '1'),
    ('    1.4 缩写说明', '2'), ('    1.5 引用文档', '2'),
    ('2. 软件设计决策', '3'),
    ('    2.1 设计目标', '3'), ('    2.2 设计原则', '3'), ('    2.3 设计约束', '3'),
    ('        2.3.1 遵循标准', '3'), ('        2.3.2 运行环境', '3'), ('        2.3.3 开发环境及工具', '3'),
    ('        2.3.4 技术限制', '4'), ('        2.3.5 其他', '4'),
    ('3. 逻辑架构设计', '5'),
    ('    3.1 设计决策', '5'), ('    3.2 软件单元', '5'),
    ('        3.2.1 界面层', '5'), ('        3.2.2 业务逻辑层', '6'), ('        3.2.3 数据访问层', '7'),
    ('    3.3 处理流程', '7'),
    ('        3.3.1 用户购票完整流程（UC_ORDER_001）', '7'),
    ('4. 人机界面设计', '9'),
    ('5. 数据存储设计', '10'),
    ('    5.1 内部数据结构', '10'), ('    5.2 数据库', '10'),
    ('        5.2.1 数据库设计', '10'), ('        5.2.2 数据表定义', '11'),
    ('    5.3 数据文件', '16'),
    ('6. 详细设计', '17'),
    ('    6.1 订单服务实现（SU_BL_05）', '17'),
    ('        6.1.1 功能描述', '17'), ('        6.1.2 处理流程', '17'), ('        6.1.3 内部数据', '18'),
    ('        6.1.4 异常与错误处理', '18'), ('        6.1.5 测试要点', '18'),
    ('    6.2 认证服务实现（SU_BL_01）', '18'),
    ('    6.3 场次服务实现（SU_BL_03）', '19'),
    ('    6.4 JWT认证组件（安全组件）', '19'),
    ('    6.5 统计服务实现（SU_BL_06）', '20'),
    ('7. 开发架构设计', '21'),
    ('    7.1 工程结构', '21'), ('    7.2 源代码文件', '22'), ('    7.3 系统组件', '22'),
    ('8. 物理架构设计', '23'),
    ('    8.1 网络环境', '23'), ('    8.2 部署方案', '23'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    # 用tab模拟目录格式
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f'{item.strip()}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    # tab对齐页码（正式编译时可在Word中自动生成目录替换此占位）
    tab_run = p.add_run(f'    {page}')
    tab_run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. 引言
# ============================================================
add_heading_styled('1. 引言', level=1)

# 1.1 编写目的
add_heading_styled('1.1 编写目的', level=2)
add_body_para('本文档用于说明TTMS电影院综合管理系统（Theater Ticket Management System）软件体系结构设计、接口设计和软件单元详细设计，是TTMS软件实现的基础。本文的预期读者包括：')
add_body_para('• 软件开发人员——理解系统架构和模块设计，指导编码实现；')
add_body_para('• 测试人员——根据详细设计编写测试用例，验证系统功能；')
add_body_para('• 项目管理人员——了解系统整体架构和技术方案，评估项目进度与风险；')
add_body_para('• 系统维护人员——掌握系统结构，便于后续维护和升级。')

# 1.2 项目概述
add_heading_styled('1.2 项目概述', level=2)
add_body_para('TTMS是一套面向中小型电影院的综合票务管理系统。系统支持顾客在线浏览影片、选座购票、支付订单、改签退票等完整购票流程，同时为影院工作人员提供影片管理、影厅管理、场次排片、订单管理、员工管理、数据统计和系统配置等后台管理功能。系统采用前后端分离架构进行开发，最终以前后端合并部署的方式运行——前端Vue 3单页应用编译为静态资源后嵌入Spring Boot应用的static目录，通过单个JAR包即可完成部署。')
add_body_para('本项目基本信息如下：')
add_body_para('• 项目名称：TTMS电影院综合管理系统（Theater Ticket Management System）；')
add_body_para('• 项目编号：XXXXXXX；')
add_body_para('• 投 资 方：指签定XXXX项目研制合同的甲方单位名称；')
add_body_para('• 用    户：指XXXX项目和软件的直接使用单位名称（电影院顾客及影院工作人员）；')
add_body_para('• 开 发 方：XXXX公司。')

# 1.3 术语定义
add_heading_styled('1.3 术语定义', level=2)
add_body_para('本文中用到的专门术语定义见表1。')
add_body_para('表1 术语定义', bold=True, first_line_indent=False)

term_rows = [
    ['1', 'JWT（JSON Web Token）', '一种无状态的认证令牌，由Header、Payload、Signature三部分组成。服务端不需要存储Session，每次请求从Token中解析用户身份信息。'],
    ['2', 'BCrypt', '一种故意慢速的密码哈希算法，自带随机盐值，每次加密结果不同，能有效抵抗暴力破解和彩虹表攻击。'],
    ['3', 'SPA（单页应用）', 'Single Page Application——只有一个HTML页面，前端通过JavaScript动态切换视图，避免整页刷新。'],
    ['4', 'ORM（对象关系映射）', '将数据库表映射为程序中的对象，使开发者可以用面向对象的方式操作数据库。'],
    ['5', '逻辑删除', '不真正删除数据库记录，而是通过deleted标记字段（0=正常/1=已删除）表示数据状态，所有查询自动过滤已删除记录。'],
    ['6', '乐观锁', '一种并发控制策略：先执行操作，如果数据已被其他事务修改则回滚重试。本项目中通过UPDATE...WHERE status=0实现座位锁定。'],
    ['7', '原子操作', '不可再分割的数据库操作，如sold_count=sold_count+N直接在SQL中运算，避免Java层"读-改-写"的并发竞态问题。'],
    ['8', '魔数校验', '通过读取文件开头的几个字节（魔数）来验证文件的真实类型，如JPEG以FF D8 FF开头，PNG以89 50 4E 47开头，防止伪造文件类型上传。'],
    ['9', 'N+1查询问题', '先执行1次查询获取N条主记录，再对每条主记录执行1次关联查询，共N+1次查询，性能随N线性增长。本系统通过批量查询优化解决。'],
    ['10', 'RESTful API', '符合REST架构风格的HTTP接口，使用GET/POST/PUT/DELETE方法对应资源的查/增/改/删操作。'],
    ['11', 'Druid', '阿里巴巴开源的数据库连接池组件，提供SQL监控、慢查询记录等运维功能。'],
    ['12', 'AOP（面向切面编程）', 'Aspect-Oriented Programming，在不修改源代码的情况下插入通用逻辑（如事务管理、日志记录）的编程技术。@Transactional即通过AOP实现。'],
    ['13', 'ThreadLocal', '线程本地存储——每个线程独立存储数据，互不干扰。Spring Security的SecurityContext基于ThreadLocal实现，请求处理完自动清理。'],
]
add_table(['序号', '术语', '含义'], term_rows, [1.5, 4, 10.5])

# 1.4 缩写说明
add_heading_styled('1.4 缩写说明', level=2)
add_body_para('本文件中用到的英文缩写说明见表2。')
add_body_para('表2 英文缩写说明', bold=True, first_line_indent=False)

abbr_rows = [
    ['1', 'TTMS', 'Theater Ticket Management System', '电影院票务管理系统'],
    ['2', 'SDS', 'Software Design Specification', '软件设计说明书'],
    ['3', 'JWT', 'JSON Web Token', 'JSON格式的Web认证令牌'],
    ['4', 'SPA', 'Single Page Application', '单页应用'],
    ['5', 'ORM', 'Object-Relational Mapping', '对象关系映射'],
    ['6', 'API', 'Application Programming Interface', '应用程序编程接口'],
    ['7', 'CRUD', 'Create, Read, Update, Delete', '增删改查操作'],
    ['8', 'CORS', 'Cross-Origin Resource Sharing', '跨域资源共享'],
    ['9', 'AOP', 'Aspect-Oriented Programming', '面向切面编程'],
    ['10', 'JDBC', 'Java Database Connectivity', 'Java数据库连接标准'],
    ['11', 'JSON', 'JavaScript Object Notation', '一种轻量级数据交换格式'],
    ['12', 'HTTP', 'HyperText Transfer Protocol', '超文本传输协议'],
    ['13', 'SQL', 'Structured Query Language', '结构化查询语言'],
    ['14', 'JDK', 'Java Development Kit', 'Java开发工具包'],
    ['15', 'CSRF', 'Cross-Site Request Forgery', '跨站请求伪造'],
]
add_table(['序号', '缩写', '原文', '说明'], abbr_rows, [1.5, 2.5, 5.5, 6.5])

# 1.5 引用文档
add_heading_styled('1.5 引用文档', level=2)
add_body_para('本文引用的文档及标准参见表3。')
add_body_para('表3 引用文档', bold=True, first_line_indent=False)

ref_rows = [
    ['1', '—', 'Spring Boot Reference Documentation', '3.2.5', '2024', 'VMware'],
    ['2', '—', 'Spring Security Reference', '6.2.4', '2024', 'VMware'],
    ['3', '—', 'MyBatis-Plus Documentation', '3.5.6', '2024', 'baomidou团队'],
    ['4', '—', 'Vue.js Guide', '3.x', '2024', 'Vue.js Core Team'],
    ['5', '—', 'MySQL 8.0 Reference Manual', '8.0', '2024', 'Oracle'],
    ['6', '—', 'JJWT Library Documentation', '0.12.5', '2024', 'JWT.io'],
    ['7', '—', 'RESTful API 设计规范', '—', '—', '—'],
]
add_table(['序号', '文档编号', '标题', '版本号', '修订日期', '编制单位'], ref_rows, [1.5, 2, 5.5, 2, 2, 3])

add_body_para('填表说明：')
add_body_para('编号：指引用文档的文件标识号，如：XUPTSE.2016.TTMS.SRS等（注：勿写成序号）；')
add_body_para('标题：指引用文档的文件名称；')
add_body_para('修订版本：指引用文档的文件版本；')
add_body_para('修订日期：指引用文档的发布日期；')
add_body_para('编写单位：指编写该引用文档的单位名称。')

doc.add_page_break()

# ============================================================
# 2. 软件设计决策
# ============================================================
add_heading_styled('2. 软件设计决策', level=1)

# 2.1 设计目标
add_heading_styled('2.1 设计目标', level=2)
add_body_para('本设计的目标如下：')
add_body_para('（1）实现用户需求：覆盖顾客端（注册登录、浏览影片、选座购票、支付、改签退票、订单管理）和管理端（影片管理、影厅管理、场次排片、订单管理、员工管理、数据统计、系统设置）的完整业务功能。')
add_body_para('（2）具有良好的可扩充性：采用分层架构设计，各层职责明确，接口清晰，新增功能时对现有代码影响最小。')
add_body_para('（3）保障数据安全与并发正确性：通过Spring Security+JWT实现认证授权，BCrypt加密存储密码，文件上传魔数校验，座位乐观锁和sold_count原子操作确保并发购票场景下数据不超卖。')
add_body_para('（4）简化部署运维：前后端合并部署于单个JAR包，启动即可运行完整系统，降低运维复杂度。')

# 2.2 设计原则
add_heading_styled('2.2 设计原则', level=2)
add_body_para('为实现上述目标，在设计软件过程中遵循以下原则：')
add_body_para('（1）分层架构原则：严格遵循"表现层（Controller）→业务逻辑层（Service）→数据访问层（Mapper）"的三层架构，上层依赖下层，下层不感知上层。')
add_body_para('（2）单一职责原则：每个类专注于单一的职责领域——AuthService仅处理认证逻辑，OrderService仅处理订单逻辑，ScheduleService仅处理场次逻辑。')
add_body_para('（3）接口隔离原则：Service层定义接口（interface），实现类（impl）与接口分离，便于扩展和测试。')
add_body_para('（4）开闭原则：通过DTO对象隔离前端请求/响应与数据库实体的差异，当需求变化时只需修改DTO而不影响核心实体。')
add_body_para('（5）防御式编程：在每个关键业务步骤进行前置校验（如订单状态校验、座位状态校验、场次时间校验），使用BusinessException及早中断异常流程。')
add_body_para('（6）最小权限原则：通过角色编码（ROLE_SUPER_ADMIN/ROLE_STAFF/ROLE_USER）和权限JSON数组实现细粒度权限控制，超级管理员可管理员工，普通员工不可操作员工管理。')
add_body_para('（7）统一响应格式：所有API返回统一的ApiResponse结构{code, message, data}，前端据此统一处理成功和错误情况。')
add_body_para('（8）优先使用数据库能力：统计聚合、原子增减等操作优先在SQL层面完成以减少数据传输和Java层的计算开销。')

# 2.3 设计约束
add_heading_styled('2.3 设计约束', level=2)

# 2.3.1
add_heading_styled('2.3.1 遵循标准', level=3)
add_body_para('本软件遵循以下标准和规范：')
add_body_para('• Java编码规范（参考阿里巴巴Java开发手册）；')
add_body_para('• RESTful API设计规范——使用标准HTTP方法（GET/POST/PUT/DELETE），URL使用名词复数形式；')
add_body_para('• JSON数据交换格式（RFC 8259）；')
add_body_para('• JWT令牌规范（RFC 7519）；')
add_body_para('• BCrypt密码加密标准；')
add_body_para('• HTTP/1.1协议（RFC 7230-7235）；')
add_body_para('• UTF-8字符编码统一使用；')
add_body_para('• 数据库命名：表名使用下划线小写，字段名使用下划线小写，Java实体使用驼峰命名。')

# 2.3.2
add_heading_styled('2.3.2 运行环境', level=3)
add_body_para('本软件系统运行的软件、硬件平台如下：')
add_body_para('（1）服务器端运行环境：')
add_body_para('• 操作系统：Windows / Linux / macOS（跨平台Java应用）；')
add_body_para('• Java运行环境：JDK 17及以上版本；')
add_body_para('• 数据库：MySQL 8.0及以上版本；')
add_body_para('• 内存：建议≥2GB（JVM堆内存）；')
add_body_para('• 存储：建议≥20GB可用空间（含数据库及上传文件）。')
add_body_para('（2）客户端运行环境：')
add_body_para('• 浏览器：现代浏览器（Chrome 90+, Firefox 90+, Edge 90+, Safari 14+），支持ES2020；')
add_body_para('• 分辨率：建议≥1366×768。')

# 2.3.3
add_heading_styled('2.3.3 开发环境及工具', level=3)
add_body_para('本软件的开发语言、环境及用到的辅助工具如下：')
add_body_para('• 后端开发语言：Java 17；')
add_body_para('• 后端框架：Spring Boot 3.2.5；')
add_body_para('• 安全框架：Spring Security 6.2.4 + JWT（jjwt 0.12.5）；')
add_body_para('• ORM框架：MyBatis-Plus 3.5.6；')
add_body_para('• 数据库连接池：Druid 1.2.22；')
add_body_para('• 前端框架：Vue 3（Composition API + Vite构建）；')
add_body_para('• HTTP客户端（前端）：Axios；')
add_body_para('• 前端状态管理：Pinia；')
add_body_para('• 前端路由：Vue Router 4（History模式）；')
add_body_para('• 前端UI框架：Element Plus；')
add_body_para('• 密码加密：Spring Security BCryptPasswordEncoder；')
add_body_para('• Excel导出：Apache POI 5.2.5；')
add_body_para('• 简化代码：Lombok；')
add_body_para('• 构建工具：Maven（后端），Vite（前端）；')
add_body_para('• 版本控制：Git；')
add_body_para('• 开发IDE：IntelliJ IDEA / VS Code；')
add_body_para('• API测试：Postman / 浏览器开发者工具。')

# 2.3.4
add_heading_styled('2.3.4 技术限制', level=3)
add_body_para('本软件在存储容量、性能、灵活性和配置等方面的约束如下：')
add_body_para('（1）性能约束：')
add_body_para('• 分页查询单页最大500条记录，防止恶意请求拖垮数据库；')
add_body_para('• 文件上传限制最大10MB，防止存储滥用；')
add_body_para('• 数据库连接池最大20个连接；Druid初始5个、最小5个连接。')
add_body_para('（2）数据约束：')
add_body_para('• 所有主表采用逻辑删除（deleted字段），数据不物理删除；')
add_body_para('• 用户名、影厅名、订单号、工号等使用UNIQUE约束保证唯一性；')
add_body_para('• 密码字段使用@JsonProperty(access=WRITE_ONLY)确保不在API响应中泄露。')
add_body_para('（3）安全约束：')
add_body_para('• JWT Token有效期为24小时，过期后需重新登录；')
add_body_para('• 未支付订单15分钟自动超时取消；')
add_body_para('• 不允许操作他人订单（只能操作自己的订单）。')

# 2.3.5
add_heading_styled('2.3.5 其他', level=3)
add_body_para('其他有关的设计考虑：')
add_body_para('（1）前后端合并部署：Vue前端编译后的静态资源放置在Spring Boot的src/main/resources/static/目录下，部署时仅需一个JAR包即可运行完整系统。WebMvcConfig配置SPA路由回退机制，确保Vue Router的History模式下刷新页面不会出现404错误。')
add_body_para('（2）数据库自动初始化：应用启动时，schema.sql自动建表（使用IF NOT EXISTS保证重复执行安全），DatabaseInitializer组件（实现CommandLineRunner）自动检查并插入默认角色、超级管理员账号、系统配置项和25部示例影片数据。')
add_body_para('（3）跨域支持：开发环境下允许http://localhost:*的跨域请求（CorsConfig.java），生产环境需根据实际域名调整CORS配置。')
add_body_para('（4）时区处理：全局统一使用Asia/Shanghai时区，日期格式统一为yyyy-MM-dd HH:mm:ss，Jackson序列化时null值不输出（减少响应体积）。')

doc.add_page_break()

# ============================================================
# 3. 逻辑架构设计
# ============================================================
add_heading_styled('3. 逻辑架构设计', level=1)
add_body_para('本系统的逻辑架构采用经典的分层架构模式，自上而下分为表现层（Controller）、业务逻辑层（Service）、数据访问层（Mapper），各层之间通过接口解耦，自上而下单向依赖。')

# 3.1 设计决策
add_heading_styled('3.1 设计决策', level=2)
add_body_para('本软件逻辑架构设计的宏观决策如下：')
add_body_para('（1）采用三层架构（3-Tier Architecture），将系统划分为表现层、业务逻辑层和数据访问层。表现层仅负责接收HTTP请求和返回JSON响应；业务逻辑层承载核心业务规则；数据访问层封装数据库操作。')
add_body_para('（2）层间通信遵循自上而下的依赖方向：Controller→Service接口→ServiceImpl→Mapper。下层对上层透明，Service层不知道HTTP请求的存在，Mapper层不知道业务规则的存在。')
add_body_para('（3）横切关注点（认证、授权、异常处理、事务管理、定时任务）通过Spring AOP机制和框架过滤器统一处理，避免在业务代码中散落。')
add_body_para('（4）DTO对象用于隔离前端请求/响应与数据库实体（Entity）之间的差异，避免实体类的序列化敏感信息泄露（如密码字段标记@JsonProperty(access=WRITE_ONLY)）。')
add_body_para('（5）认证授权采用Spring Security过滤器链+JWT无状态令牌的方案，白名单路径跳过认证，管理端路径需管理员角色。')

add_body_para('图1  TTMS软件逻辑架构', bold=True, first_line_indent=False)
add_body_para('（架构图——待补充）', first_line_indent=False)
add_body_para('┌──────────────────────────────────────────────────────────────────┐')
add_body_para('│                     表 现 层 (Controller)                         │')
add_body_para('│  AuthController | MovieController | ScheduleController             │')
add_body_para('│  OrderController | HallController | EmployeeController            │')
add_body_para('│  AdminOrderController | StatisticsController | SystemController   │')
add_body_para('│  FileController                                                    │')
add_body_para('├──────────────────────────────────────────────────────────────────┤')
add_body_para('│                     DTO (数据传输对象)                             │')
add_body_para('│  ApiResponse | LoginRequest/Response | OrderRequest | ...         │')
add_body_para('├──────────────────────────────────────────────────────────────────┤')
add_body_para('│                  业 务 逻 辑 层 (Service)                          │')
add_body_para('│  AuthService | MovieService | ScheduleService | OrderService      │')
add_body_para('│  HallService | StatisticsService                                   │')
add_body_para('├──────────────────────────────────────────────────────────────────┤')
add_body_para('│                   安 全 过 滤 器 链                                │')
add_body_para('│  JwtAuthenticationFilter | SecurityConfig | JwtTokenProvider       │')
add_body_para('├──────────────────────────────────────────────────────────────────┤')
add_body_para('│                  数 据 访 问 层 (Mapper)                           │')
add_body_para('│  UserMapper | EmployeeMapper | MovieMapper | HallMapper           │')
add_body_para('│  ScheduleMapper | SeatMapper | OrderMapper | OrderLogMapper       │')
add_body_para('│  RoleMapper | SystemConfigMapper                                   │')
add_body_para('├──────────────────────────────────────────────────────────────────┤')
add_body_para('│                     数 据 库 (MySQL 8.0)                          │')
add_body_para('│               10张业务表 + InnoDB引擎 + utf8mb4编码               │')
add_body_para('└──────────────────────────────────────────────────────────────────┘')

# 3.2 软件单元
add_heading_styled('3.2 软件单元', level=2)
add_body_para('本小节给出各层类的构成及其关系。')

# 3.2.1
add_heading_styled('3.2.1 界面层', level=3)
add_body_para('界面层由10个Controller类构成，负责接收HTTP请求、参数校验、调用业务层、返回JSON响应。各Controller均使用@RestController注解，通过@RequestMapping定义URL前缀。')
add_body_para('图2 界面层类图 （待补充）', bold=True, first_line_indent=False)
add_body_para('界面层的软件单元构成如表4所示。')
add_body_para('表4 界面层软件单元构成', bold=True, first_line_indent=False)

ui_rows = [
    ['1', 'SU_UI_01', 'AuthController', '用户注册、登录、修改密码', '/api/auth'],
    ['2', 'SU_UI_02', 'MovieController', '影片列表/搜索/详情(公开)；增删改/设置热门/设置状态(管理端)', '/api/movies'],
    ['3', 'SU_UI_03', 'ScheduleController', '场次查询/座位矩阵查看(公开)；场次CRUD(管理端)', '/api/schedules'],
    ['4', 'SU_UI_04', 'OrderController', '用户端订单：创建、支付、改签、退票、查看我的订单', '/api/user/orders'],
    ['5', 'SU_UI_05', 'AdminOrderController', '管理端订单：全部订单列表、订单详情、协助下单', '/api/admin/orders'],
    ['6', 'SU_UI_06', 'HallController', '影厅CRUD管理、状态设置', '/api/admin/halls'],
    ['7', 'SU_UI_07', 'EmployeeController', '员工管理：列表、增删改查、重置密码、启用/禁用', '/api/admin/employees'],
    ['8', 'SU_UI_08', 'StatisticsController', '营收概览、每日/月度趋势、影片排行、Excel报表导出', '/api/admin/statistics'],
    ['9', 'SU_UI_09', 'SystemController', '系统配置管理、操作日志查看、用户主题偏好', '/api/admin/system'],
    ['10', 'SU_UI_10', 'FileController', '图片文件上传（含魔数校验）', '/api/upload'],
]
add_table(['序号', '软件单元标识符', '软件单元（类）名称', '功能说明', '备注（URL前缀）'], ui_rows, [1, 3, 3.5, 5.5, 3])

# 3.2.2
add_heading_styled('3.2.2 业务逻辑层', level=3)
add_body_para('业务逻辑层由6个Service接口及对应的实现类构成，承载系统的核心业务规则。每个Service封装一个业务领域的所有操作逻辑。')
add_body_para('业务逻辑层的软件单元构成如表5所示。')
add_body_para('表5 业务逻辑层软件单元构成', bold=True, first_line_indent=False)

svc_rows = [
    ['1', 'SU_BL_01', 'AuthService', 'AuthServiceImpl', '认证逻辑：登录（区分USER/ADMIN登录类型）、注册（唯一性校验+BCrypt加密）、修改密码'],
    ['2', 'SU_BL_02', 'MovieService', 'MovieServiceImpl', '影片管理：条件分页查询、关键词搜索、CRUD、热门设置、上下架状态管理'],
    ['3', 'SU_BL_03', 'ScheduleService', 'ScheduleServiceImpl', '场次管理：CRUD、按影片/影厅查询、座位矩阵生成（延迟初始化）、时间冲突检测、场次信息批量填充（N+1优化）'],
    ['4', 'SU_BL_04', 'HallService', 'HallServiceImpl', '影厅管理：CRUD、删除前检查进行中场次、设置维护状态'],
    ['5', 'SU_BL_05', 'OrderService', 'OrderServiceImpl', '订单管理：创建订单（乐观锁座位+原子操作售出计数）、支付、改签（差价计算）、退票、超时取消定时任务（@Scheduled）'],
    ['6', 'SU_BL_06', 'StatisticsService', 'StatisticsServiceImpl', '数据统计：营收概览、每日/月度趋势、影片票房排行、Excel报表导出（Apache POI）'],
]
add_table(['序号', '软件单元标识符', '软件单元（接口）名称', '软件单元（实现类）名称', '功能说明'], svc_rows, [1, 2.5, 3, 3.5, 6])

# 3.2.3
add_heading_styled('3.2.3 数据访问层', level=3)
add_body_para('数据访问层由10个Mapper接口构成，均继承MyBatis-Plus的BaseMapper<T>，自动获得基本的CRUD方法。部分Mapper定义了自定义SQL方法用于复杂查询和批量操作。业务层通过调用Mapper接口访问数据库。')
add_body_para('数据访问层的软件单元构成如表6所示。')
add_body_para('表6 数据访问层软件单元构成', bold=True, first_line_indent=False)

mapper_rows = [
    ['1', 'SU_DA_01', 'UserMapper', 'User', 'findByUsername（含deleted=0过滤）'],
    ['2', 'SU_DA_02', 'EmployeeMapper', 'Employee', 'findByUsername（含deleted=0过滤）'],
    ['3', 'SU_DA_03', 'MovieMapper', 'Movie', 'searchMovies（关键词模糊搜索，含deleted=0过滤）'],
    ['4', 'SU_DA_04', 'HallMapper', 'Hall', '继承BaseMapper'],
    ['5', 'SU_DA_05', 'ScheduleMapper', 'Schedule', 'selectByMovieId、selectByHallId、selectUpcoming、selectByHallAndTimeRange（冲突检测）'],
    ['6', 'SU_DA_06', 'SeatMapper', 'Seat', 'lockSeat（乐观锁）、markSold（标记售出）、releaseSeatsByOrderId（批量释放）'],
    ['7', 'SU_DA_07', 'OrderMapper', 'Order', 'selectByOrderNo（查重）、selectExpiredOrders（超时扫描）、incrementSoldCount/decrementSoldCount（原子操作）'],
    ['8', 'SU_DA_08', 'OrderLogMapper', 'OrderLog', '继承BaseMapper'],
    ['9', 'SU_DA_09', 'RoleMapper', 'Role', '继承BaseMapper'],
    ['10', 'SU_DA_10', 'SystemConfigMapper', 'SystemConfig', '继承BaseMapper'],
]
add_table(['序号', '软件单元标识符', '软件单元（类）名称', '对应实体', '主要功能（含自定义方法）'], mapper_rows, [1, 2.5, 3, 2.5, 7])

# 3.3 处理流程
add_heading_styled('3.3 处理流程', level=2)
add_body_para('本节以系统最核心的用例——"用户购票"为例，给出从注册到退票的完整处理流程。')

# 3.3.1
add_heading_styled('3.3.1 用户购票完整流程（用例标识符 UC_ORDER_001）', level=3)
add_body_para('该用例覆盖顾客从注册到完成购票、退票的完整链路，涉及AuthController、MovieController、ScheduleController、OrderController四个控制器和对应的Service层。')
add_body_para('图3 用户购票用例处理流程（顺序图——待补充）', bold=True, first_line_indent=False)

add_body_para('步骤1：用户注册', bold=True, first_line_indent=False)
add_body_para('前端POST /api/auth/register→AuthController.register()→AuthServiceImpl.register()：检查username在user表和employee表中均不存在→BCrypt密码加密→设置默认status=0和theme="white"→插入user表。@Transactional保证原子性。')

add_body_para('步骤2：用户登录', bold=True, first_line_indent=False)
add_body_para('前端POST /api/auth/login→AuthController.login()→AuthServiceImpl.login()：根据loginType="USER"查user表→BCrypt.matches()验证密码→JwtTokenProvider.generateToken()生成JWT Token（载荷含userId/username/role/loginType，有效期24小时，HMAC-SHA256签名）→返回LoginResponse{token, userId, username, roleCode, roleName, permissions, theme}。')

add_body_para('步骤3：浏览影片列表', bold=True, first_line_indent=False)
add_body_para('前端GET /api/movies/list?page=1&size=50&status=1→MovieController.list()→MovieServiceImpl.list()：MyBatis-Plus分页插件自动生成COUNT+SELECT LIMIT SQL，自动追加deleted=0过滤，按sort_order DESC, create_time DESC排序返回。')

add_body_para('步骤4：查看场次和座位', bold=True, first_line_indent=False)
add_body_para('前端GET /api/schedules/query/movie/{movieId}→ScheduleController.queryByMovie()→ScheduleServiceImpl.queryByMovie()：LEFT JOIN movie和hall表返回含影片名、影厅名、可用座位数的场次列表。前端GET /api/schedules/query/{scheduleId}/seats→ScheduleController.getSeats()→ScheduleServiceImpl.getSeats()：查seat表→若为空则根据Hall的rowCount×colCount首次自动生成座位→parse hall.seat_layout标记不可用座位(status=3)→转二维矩阵返回。')

add_body_para('步骤5：下单（核心事务）', bold=True, first_line_indent=False)
add_body_para('前端POST /api/user/orders/create→OrderController.createOrder()→OrderServiceImpl.createOrder()——整个方法在@Transactional保护下执行，共8步：')
add_body_para('①验证场次存在且status=1(正常)且未开场→②逐一检查座位存在且status=0(空闲)→③生成订单号(yyyyMMdd+8位随机字符A-Z0-9，递归查重)→④INSERT订单(status=0待支付)→⑤原子锁定座位(UPDATE seat SET status=1, lock_time=NOW(), order_id=? WHERE id=? AND status=0——WHERE status=0为乐观锁条件，若affected_rows≠1则回滚抛异常)→⑥更新订单总价(场次票价×座位数)→⑦原子更新售出数(UPDATE schedule SET sold_count=sold_count+N WHERE id=?)→⑧写操作日志(order_log, operation_type=CREATE)。')

add_body_para('步骤6：支付', bold=True, first_line_indent=False)
add_body_para('前端POST /api/user/orders/pay/{orderId}→OrderController.payOrder()→OrderServiceImpl.payOrder()：验证是本用户的订单+status=0(待支付)+场次未开场→逐一标记座位已售出(status=1→2)→更新订单status=1(待观影)+记录pay_time→写操作日志(PAY)。支付时不重复增加sold_count（下单时已预增）。')

add_body_para('步骤7：改签（可选）', bold=True, first_line_indent=False)
add_body_para('前端POST /api/user/orders/reschedule→OrderController.reschedule()→OrderServiceImpl.reschedule()——最复杂事务：验证原订单status=1且是本人订单且原场次未开场→验证新场次有效且≠原场次→验证新座位全部空闲→释放原座位(批量UPDATE status=0)→原子减少原场次sold_count→原订单status改为3(已改签)→创建新订单：计算价差，若需补差价则新订单status=0(待支付)，否则新订单status=1(直接生效)→写操作日志(RESCHEDULE)。')

add_body_para('步骤8：退票（可选）', bold=True, first_line_indent=False)
add_body_para('前端POST /api/user/orders/refund/{orderId}→OrderController.refund()→OrderServiceImpl.refund()：验证订单status=1且是本人订单且场次未开场→批量释放座位(status=0)→原子减少sold_count(WHERE sold_count>=N保护下限)→订单status改为4(已退票)→写操作日志(REFUND)。')

add_body_para('补充：超时自动取消', bold=True, first_line_indent=False)
add_body_para('@Scheduled(fixedDelay=120000)每2分钟执行cancelExpired()→查询status=0且创建超过15分钟的订单→逐一释放座位→减少sold_count→改status=5(已过期)→写操作日志(EXPIRE)。每个订单独立try-catch，单失败不影响其他。本任务天然幂等：订单status改为5后不再被查询条件匹配。')

doc.add_page_break()

# ============================================================
# 4. 人机界面设计
# ============================================================
add_heading_styled('4. 人机界面设计', level=1)
add_body_para('本系统前端采用Vue 3 + Element Plus组件库构建SPA（单页应用），界面设计遵循以下策略：')
add_body_para('（1）响应式布局：采用Element Plus的栅格系统，适配1366×768及以上分辨率，页面内容居中展示。')
add_body_para('（2）组件化设计：导航栏（NavBar）、影片卡片、座位选择器、订单列表等封装为可复用的Vue组件。')
add_body_para('（3）权限驱动UI：根据登录接口返回的permissions数组和roleCode，通过v-if/v-show控制菜单项和功能按钮的显示/隐藏。SUPER_ADMIN可见"员工管理"和"系统设置"，STAFF不可见，USER仅可见用户端页面。')
add_body_para('（4）路由守卫：Vue Router的beforeEach导航守卫检查token是否存在和角色是否匹配，未认证用户访问需登录页面时自动重定向到/login。')
add_body_para('（5）主题切换：支持日间（白色）和夜间（暗黑）两种主题，用户可在个人中心切换并保存偏好到后端。')

add_body_para('系统主要界面清单如下：', first_line_indent=True)
add_body_para('表7 系统前端页面构成', bold=True, first_line_indent=False)

page_rows = [
    ['1', '首页', '/', '热门影片展示、正在热映列表、搜索入口', '公开'],
    ['2', '登录页', '/login', '用户/管理员登录表单，含loginType选择', '公开'],
    ['3', '注册页', '/register', '新用户注册表单', '公开'],
    ['4', '影片详情页', '/movie/:id', '影片信息、场次列表、选座入口', '公开'],
    ['5', '选座页面', '/seats/:scheduleId', '影厅座位矩阵可视化选座', '需认证'],
    ['6', '我的订单', '/my-orders', '当前用户的订单列表和状态筛选', '需认证'],
    ['7', '个人中心', '/profile', '个人信息修改、主题切换、修改密码', '需认证'],
    ['8', '管理仪表盘', '/admin/dashboard', '关键数据概览卡片', '管理员'],
    ['9', '影片管理', '/admin/movies', '影片CRUD表格、上下架、热门设置', '管理员'],
    ['10', '影厅管理', '/admin/halls', '影厅CRUD、维护状态管理', '管理员'],
    ['11', '场次管理', '/admin/schedules', '排片CRUD、时间冲突自动检测提示', '管理员'],
    ['12', '订单管理', '/admin/orders', '全部订单列表、状态筛选、协助下单', '管理员'],
    ['13', '员工管理', '/admin/employees', '员工CRUD、重置密码、启用/禁用', '超级管理员'],
    ['14', '数据统计', '/admin/statistics', '营收图表、排行、Excel报表导出', '管理员'],
    ['15', '系统设置', '/admin/settings', '影院名称、联系电话、订单超时时间、公告', '超级管理员'],
]
add_table(['序号', '页面名称', '对应路由', '功能说明', '访问权限'], page_rows, [1, 3, 4, 6, 2])

add_body_para('前端技术栈补充：Axios封装（请求拦截器自动添加Authorization Header，响应拦截器统一处理401跳转登录页）；Pinia Store存储全局状态（用户信息、token、角色权限、主题偏好）；Vue Router History模式（后端WebMvcConfig配置SPA路由回退，非API非静态资源请求全部返回index.html）；Element Plus组件库提供一致UI风格。')

doc.add_page_break()

# ============================================================
# 5. 数据存储设计
# ============================================================
add_heading_styled('5. 数据存储设计', level=1)

# 5.1 内部数据结构
add_heading_styled('5.1 内部数据结构', level=2)
add_body_para('本系统的内存数据结构主要包括：')
add_body_para('（1）SecurityContext上下文：每个HTTP请求处理期间，JwtAuthenticationFilter将userId（作为Principal）和role（作为GrantedAuthority）存入SecurityContextHolder。Controller通过SecurityContextHolder.getContext().getAuthentication()获取当前用户信息。上下文基于ThreadLocal实现，请求处理完成后自动清理。')
add_body_para('（2）MyBatis-Plus分页对象：Page<T>承载分页参数（current页码，size页大小）和结果（records记录列表，total总条数）。分页插件自动设置最大单页500条。')
add_body_para('（3）ApiResponse统一响应对象：{Integer code, String message, Object data}，通过静态工厂方法success()和error()快速构建。data为泛型，可以是单个对象、列表、Map或null。')
add_body_para('（4）DTO对象（请求/响应传输）：LoginRequest{username, password, loginType}、LoginResponse{token, userId, username, realName, roleCode, roleName, permissions, theme}、OrderRequest{scheduleId, seatNumbers[], userId?}、RegisterRequest{username, password, phone, email, nickname}、RescheduleRequest{orderId, newScheduleId, newSeatNumbers[]}。以上DTO均为请求/响应期间临时对象，不持久化。')
add_body_para('（5）全局变量：本系统不定义Java层面的全局变量。系统配置（影院名称、联系电话、订单超时时间等）存储在system_config表中，通过SystemController动态读取和更新。角色权限定义在DatabaseInitializer常量中。')

# 5.2 数据库
add_heading_styled('5.2 数据库', level=2)

# 5.2.1
add_heading_styled('5.2.1 数据库设计', level=3)
add_body_para('本系统使用MySQL 8.0关系型数据库，数据库名为TTMS（可通过JDBC URL的createDatabaseIfNotExist=true自动创建），采用InnoDB存储引擎和utf8mb4字符集。数据库包含10张业务表。')
add_body_para('数据库表关系如图4所示（ER图——待补充）：', bold=True, first_line_indent=False)

add_body_para('''
                    ┌──────────┐            ┌──────────┐
                    │   role   │            │  user    │
                    │ (角色表)  │            │ (用户表)  │
                    └────┬─────┘            └────┬─────┘
                         │role_id                │user_id
                    ┌────▼─────┐            ┌────▼──────────────────────┐
                    │ employee │            │          order            │
                    │ (员工表)  │            │         (订单表)           │
                    └──────────┘            │ movie_id(冗余) hall_id(冗余)│
                                           │ schedule_id                │
┌──────────┐      ┌──────────┐             └──┬────┬────┬────┬────────┘
│  movie   │      │   hall   │          order_id│    │    │original
│ (影片表)  │      │ (影厅表)  │          ┌──────▼┐  │    │_order_id
└────┬─────┘      └────┬─────┘          │  seat │  │    │
     │movie_id         │hall_id          │ (座位表)│  │    │
     │           ┌─────▼──────┐          └───────┘  │    │
     └──────────►│  schedule  │              ┌──────▼──────▼──┐
                 │  (场次表)   │              │   order_log    │
                 └────────────┘              │  (操作日志表)    │
                                            └────────────────┘
                 ┌──────────────┐
                 │system_config │
                 │ (系统配置表)  │
                 └──────────────┘
''', first_line_indent=False)

add_body_para('核心关系说明：')
add_body_para('• role 1:N employee——一个角色（如STAFF）可关联多名员工；')
add_body_para('• movie 1:N schedule——一部影片可排多个场次；')
add_body_para('• hall 1:N schedule——一个影厅可排多个场次；')
add_body_para('• schedule 1:N seat——一个场次包含多个座位（按rowCount×colCount生成）；')
add_body_para('• user 1:N order——一个用户可有多个订单；')
add_body_para('• schedule 1:N order——一个场次可产生多个订单；')
add_body_para('• order 1:N seat——一个订单可包含多个座位；')
add_body_para('• order 1:N order_log——一个订单可有多条操作日志（CREATE/PAY/RESCHEDULE/REFUND/EXPIRE）；')
add_body_para('• order冗余movie_id和hall_id——减少多表JOIN，提升查询性能，同时保留历史快照。')

# 5.2.2
add_heading_styled('5.2.2 数据表定义', level=3)
add_body_para('本系统共10张数据表，详细定义如下。所有表均使用InnoDB引擎和utf8mb4字符集，核心业务表均包含create_time（自动填充创建时间）、update_time（自动填充更新时间）、deleted（逻辑删除标记，0=未删除，1=已删除）三个通用字段。')

# role
add_body_para('表8 role（角色表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '角色ID'],
    ['role_code', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '角色编码：ROLE_SUPER_ADMIN/ROLE_STAFF/ROLE_USER'],
    ['role_name', 'VARCHAR(50)', 'NOT NULL', '角色名称：超级管理员/普通员工/普通用户'],
    ['description', 'VARCHAR(255)', '', '角色描述'],
    ['permissions', 'TEXT', '', '权限JSON数组，控制前端功能按钮显示'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除标记'],
], [3, 4, 4, 5])

# user
add_body_para('表9 user（用户表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '用户ID'],
    ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '用户名'],
    ['password', 'VARCHAR(255)', 'NOT NULL', '密码（BCrypt加密存储，@JsonProperty WRITE_ONLY防止泄露）'],
    ['phone', 'VARCHAR(20)', '', '手机号'],
    ['email', 'VARCHAR(100)', '', '邮箱'],
    ['nickname', 'VARCHAR(50)', '', '昵称'],
    ['avatar', 'VARCHAR(255)', '', '头像URL'],
    ['status', 'INT', 'DEFAULT 0', '账号状态：0-正常，1-禁用'],
    ['theme', 'VARCHAR(20)', "DEFAULT 'white'", '用户主题偏好：white（日间）/ dark（暗黑）'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除'],
], [3, 3.5, 4.5, 5])

# employee
add_body_para('表10 employee（员工表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '员工ID'],
    ['employee_no', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '工号'],
    ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '用户名'],
    ['password', 'VARCHAR(255)', 'NOT NULL', '密码（BCrypt加密，@JsonProperty WRITE_ONLY）'],
    ['real_name', 'VARCHAR(50)', '', '真实姓名'],
    ['phone', 'VARCHAR(20)', '', '手机号'],
    ['role_id', 'BIGINT', '', '关联角色ID（关联role表）'],
    ['status', 'INT', 'DEFAULT 0', '账号状态：0-正常，1-禁用'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除'],
], [3, 3.5, 4.5, 5])

# movie
add_body_para('表11 movie（影片表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '影片ID'],
    ['movie_name', 'VARCHAR(100)', 'NOT NULL', '片名'],
    ['genre', 'VARCHAR(100)', '', '类型（逗号分隔多个标签，如"科幻,冒险"）'],
    ['duration', 'INT', '', '时长（分钟）'],
    ['actors', 'VARCHAR(500)', '', '主演（逗号分隔）'],
    ['director', 'VARCHAR(100)', '', '导演'],
    ['description', 'TEXT', '', '影片简介'],
    ['poster_url', 'VARCHAR(500)', '', '海报图片URL'],
    ['release_date', 'DATE', '', '上映日期'],
    ['base_price', 'DECIMAL(10,2)', '', '基础票价（元）'],
    ['status', 'INT', 'DEFAULT 1', '状态：0-下架，1-上架(热映)，2-即将上映'],
    ['is_hot', 'INT', 'DEFAULT 0', '是否热门：0-否，1-是（置顶显示）'],
    ['sort_order', 'INT', 'DEFAULT 0', '排序权重（越大越靠前）'],
    ['country', 'VARCHAR(50)', '', '国家/地区'],
    ['language', 'VARCHAR(50)', '', '语言'],
    ['rating', 'DOUBLE', 'DEFAULT 0', '评分（1.0-10.0）'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除'],
], [3, 3, 5, 5])

# hall
add_body_para('表12 hall（影厅表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '影厅ID'],
    ['hall_name', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '影厅名称（如"1号标准厅"）'],
    ['row_count', 'INT', 'NOT NULL', '座位行数'],
    ['col_count', 'INT', 'NOT NULL', '座位列数'],
    ['capacity', 'INT', '', '总容量（行数×列数-不可用座位数）'],
    ['hall_type', 'VARCHAR(20)', "DEFAULT 'STANDARD'", '影厅类型：STANDARD（普通）/IMAX（巨幕）/VIP（豪华）/4DX'],
    ['status', 'INT', 'DEFAULT 1', '状态：0-维护中，1-正常'],
    ['remark', 'VARCHAR(255)', '', '备注'],
    ['seat_layout', 'TEXT', '', '不可用座位位置JSON数组，如["1-5","2-10"]表示第1行第5列和第2行第10列不可用'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除'],
], [3, 3.5, 4.5, 5])

# schedule
add_body_para('表13 schedule（场次表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '场次ID'],
    ['movie_id', 'BIGINT', 'NOT NULL', '关联影片ID'],
    ['hall_id', 'BIGINT', 'NOT NULL', '关联影厅ID'],
    ['start_time', 'DATETIME', 'NOT NULL', '放映开始时间'],
    ['end_time', 'DATETIME', '', '放映结束时间（start_time + movie.duration）'],
    ['price', 'DECIMAL(10,2)', '', '实际票价（可覆盖影片base_price）'],
    ['status', 'INT', 'DEFAULT 1', '状态：0-已取消，1-正常放映，2-已结束'],
    ['sold_count', 'INT', 'DEFAULT 0', '已售/已锁定座位数（用于计算剩余可用座位），通过原子操作增减'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除'],
], [3, 3, 5, 5])

# seat
add_body_para('表14 seat（座位表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '座位ID'],
    ['schedule_id', 'BIGINT', 'NOT NULL', '关联场次ID'],
    ['seat_row', 'INT', 'NOT NULL', '座位行号（从1开始）'],
    ['seat_col', 'INT', 'NOT NULL', '座位列号（从1开始）'],
    ['seat_number', 'VARCHAR(20)', 'NOT NULL', '座位编号（如"A-05"，行号用字母表示）'],
    ['status', 'INT', 'DEFAULT 0', '状态：0-空闲，1-已锁定(待支付)，2-已售出，3-过道/不可用'],
    ['lock_time', 'DATETIME', '', '锁定时间（配合定时任务判断超时）'],
    ['order_id', 'BIGINT', '', '关联订单ID'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
], [3, 3, 5, 5])

# order
add_body_para('表15 order（订单表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '订单ID'],
    ['order_no', 'VARCHAR(30)', 'NOT NULL, UNIQUE', '订单号（格式：yyyyMMdd+8位随机字符A-Z0-9）'],
    ['user_id', 'BIGINT', 'NOT NULL', '下单用户ID'],
    ['schedule_id', 'BIGINT', 'NOT NULL', '关联场次ID'],
    ['movie_id', 'BIGINT', '', '冗余影片ID（减少JOIN，保留历史快照）'],
    ['hall_id', 'BIGINT', '', '冗余影厅ID（减少JOIN）'],
    ['seat_numbers', 'VARCHAR(500)', '', '座位编号列表（逗号分隔，如"A-05,A-06"）'],
    ['seat_count', 'INT', 'DEFAULT 1', '座位数量'],
    ['total_price', 'DECIMAL(10,2)', '', '订单总金额（票价×座位数，改签时可能含差价）'],
    ['status', 'INT', 'DEFAULT 0', '状态：0-待支付，1-待观影，2-已完成，3-已改签，4-已退票，5-已过期'],
    ['pay_time', 'DATETIME', '', '支付时间'],
    ['original_order_id', 'BIGINT', '', '原订单ID（改签来源，形成改签链路追踪）'],
    ['reschedule_time', 'DATETIME', '', '改签/退票时间'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
], [3, 3.5, 5, 4.5])

# order_log
add_body_para('表16 order_log（操作日志表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '日志ID'],
    ['order_id', 'BIGINT', '', '关联订单ID'],
    ['operation_type', 'VARCHAR(20)', '', '操作类型：CREATE/PAY/RESCHEDULE/REFUND/EXPIRE'],
    ['before_content', 'TEXT', '', '操作前订单内容（JSON）'],
    ['after_content', 'TEXT', '', '操作后订单内容（JSON）'],
    ['operator_id', 'BIGINT', '', '操作人ID'],
    ['operator_type', 'VARCHAR(20)', '', '操作人类型：USER(顾客)/EMPLOYEE(员工)/SYSTEM(系统定时任务)'],
    ['remark', 'VARCHAR(500)', '', '操作备注描述'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
], [3, 3, 5, 5])

# system_config
add_body_para('表17 system_config（系统配置表）', bold=True, first_line_indent=False)
add_table(['字段名', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PRIMARY KEY AUTO_INCREMENT', '配置ID'],
    ['config_key', 'VARCHAR(100)', 'NOT NULL, UNIQUE', '配置键：theme/cinema_name/contact_phone/order_timeout/notice'],
    ['config_value', 'TEXT', '', '配置值'],
    ['description', 'VARCHAR(255)', '', '配置项描述'],
    ['create_time', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
    ['update_time', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '更新时间'],
    ['deleted', 'INT', 'DEFAULT 0', '逻辑删除'],
], [3, 3.5, 4.5, 5])

# 5.3 数据文件
add_heading_styled('5.3 数据文件', level=2)
add_body_para('本系统的数据文件主要包括：')
add_body_para('（1）配置文件（application.yml）：位于src/main/resources/目录，包含服务器端口（8080）、数据库连接信息（MySQL连接URL/用户名/密码/Druid连接池参数）、JWT密钥和过期时间（24小时=86400000ms）、文件上传参数（最大10MB，存储路径./uploads/）、Jackson序列化配置（时区Asia/Shanghai、日期格式yyyy-MM-dd HH:mm:ss、null值不序列化）、MyBatis-Plus配置（驼峰转下划线、逻辑删除字段deleted、ID自增）等。')
add_body_para('（2）建表脚本（schema.sql）：位于src/main/resources/目录，包含10张表的DDL语句（全部使用CREATE TABLE IF NOT EXISTS保证重复执行安全），以及ALTER TABLE语句（兼容旧数据库增量升级）。还包括3个示例影厅的INSERT...WHERE NOT EXISTS语句。配置spring.sql.init.mode=always确保每次启动自动执行，continue-on-error=true确保增量执行时兼容已存在的列。')
add_body_para('（3）上传文件：影片海报等图片上传至./uploads/目录下按日期组织的子目录中。上传文件经过Content-Type和魔数双重校验（检查JPEG FF D8 FF、PNG 89 50 4E 47、GIF 47 49 46、BMP 42 4D、WebP 52 49 46 46等文件头魔数），仅允许真实图片格式。')
add_body_para('（4）日志文件：Spring Boot默认日志输出到控制台（stdout），MyBatis-Plus SQL日志通过StdOutImpl输出到控制台。生产环境可配置logback-spring.xml实现文件滚动策略。')

doc.add_page_break()

# ============================================================
# 6. 详细设计
# ============================================================
add_heading_styled('6. 详细设计', level=1)
add_body_para('本节给出每个核心软件单元的具体设计细节。')

# 6.1 OrderServiceImpl
add_heading_styled('6.1 软件单元1（SU_BL_05: OrderServiceImpl——订单服务实现）', level=2)

add_heading_styled('6.1.1 功能描述', level=3)
add_body_para('OrderServiceImpl是本系统最核心、最复杂的业务逻辑组件，实现了OrderService接口，负责处理从下单到退票的完整订单生命周期。该类使用@RequiredArgsConstructor注入OrderMapper、SeatMapper、ScheduleMapper、OrderLogMapper等依赖，核心方法均标注@Transactional确保数据一致性。')
add_body_para('依赖关系：依赖OrderMapper（订单CRUD及超时扫描）→依赖SeatMapper（座位锁定/标记售出/批量释放）→依赖ScheduleMapper（场次信息查询及sold_count原子增减）→依赖OrderLogMapper（操作日志记录）。被OrderController（用户端）和AdminOrderController（管理端）调用，暴露REST API接口。此外含@Scheduled方法cancelExpired()作为独立定时任务，无需外部调用。')

add_heading_styled('6.1.2 处理流程', level=3)
add_body_para('（1）createOrder(OrderRequest, Long userId)——创建订单', bold=True, first_line_indent=False)
add_body_para('流程：验证场次存在且status=1(正常)且startTime>now(未开场)→遍历seatNumbers逐一检查座位存在且status=0(空闲)→调用generateOrderNo()生成唯一订单号（yyyyMMdd+8位随机字符，递归查重直到唯一）→INSERT订单记录(status=0待支付)→调用seatMapper.lockSeat()原子锁定每个座位（SQL: UPDATE seat SET status=1, lock_time=NOW(), order_id=? WHERE id=? AND status=0，WHERE status=0为乐观锁条件，affected_rows≠1则抛异常回滚）→计算总价=场次票价×座位数并UPDATE订单→调用scheduleMapper.incrementSoldCount()原子增加场次售出数（SQL: UPDATE schedule SET sold_count=sold_count+N WHERE id=?）→调用orderLogMapper.insert()记录CREATE日志。关键并发控制：乐观锁WHERE status=0确保两个并发请求同一座位时只有一个成功。')

add_body_para('（2）payOrder(Long orderId, Long userId)——支付订单', bold=True, first_line_indent=False)
add_body_para('流程：查询订单→验证是本用户的订单(userId==order.userId)→验证status=0(待支付)→验证场次未开场(startTime>now)→逐一markSold将座位status从1(已锁定)改为2(已售出)→UPDATE订单status=1(待观影)+记录pay_time=NOW()→记录PAY日志。注意：支付时不重复增加sold_count，因为下单时已预先增加（见设计决策说明）。')

add_body_para('（3）reschedule(RescheduleRequest, Long userId)——改签', bold=True, first_line_indent=False)
add_body_para('最复杂业务流程（约80行代码）：验证原订单status=1(待观影)且是本人订单且原场次未开场→验证新场次存在有效且不同于原场次(不能改签到同一场次)→验证新座位全部status=0(空闲)→releaseSeatsByOrderId批量释放原座位(UPDATE seat SET status=0, lock_time=NULL, order_id=NULL WHERE order_id=?)→原子减少原场次sold_count→原订单status改为3(已改签)→创建新订单：计算价差(新总价-旧总价)，若需补差价(新>旧)则新订单status=0(待支付)，否则新订单status=1(直接生效)且新座位直接标记为已售出→记录RESCHEDULE日志(含价差信息)。')

add_body_para('（4）refund(Long orderId, Long userId)——退票', bold=True, first_line_indent=False)
add_body_para('流程：查询订单→验证status=1(待观影)且是本人订单且场次未开场→releaseSeatsByOrderId批量释放座位→decrementSoldCount原子减少场次售出数(UPDATE schedule SET sold_count=sold_count-N WHERE id=? AND sold_count>=N，有下限保护)→订单status改为4(已退票)+记录reschedule_time→记录REFUND日志。')

add_body_para('（5）cancelExpired()——超时取消定时任务', bold=True, first_line_indent=False)
add_body_para('@Scheduled(fixedDelay=120000)每2分钟自动执行。流程：selectExpiredOrders(15)查询status=0 AND create_time<15分钟前的订单→遍历每个过期订单：释放座位→减少sold_count→status改为5(已过期)→记录EXPIRE日志(operator_type=SYSTEM)。每个订单独立try-catch，单个失败不影响其他订单处理。天然幂等：订单status改为5后不再满足查询条件。')

add_heading_styled('6.1.3 内部数据', level=3)
add_body_para('该类无独立状态属性（无字段），所有数据通过方法参数传入和Mapper接口持久化。依赖的Mapper通过构造函数注入（@RequiredArgsConstructor+final字段），在方法内使用LambdaQueryWrapper构建动态查询条件。订单金额使用BigDecimal类型确保精度（避免浮点运算误差）。')

add_heading_styled('6.1.4 异常与错误处理', level=3)
add_body_para('主要异常情况及其BusinessException消息：')
add_body_para('• 场次不存在或已取消→"场次不存在或已取消"；')
add_body_para('• 场次已开场→"场次已开始放映，无法操作"；')
add_body_para('• 座位不存在→"座位XXX不存在"；')
add_body_para('• 座位已被占用→"座位XXX已被占用，可能已被其他用户抢占"（乐观锁失败触发回滚）；')
add_body_para('• 订单状态不匹配→"订单状态不正确"（如已支付的订单不能再次支付）；')
add_body_para('• 无权操作他人订单→"无权操作他人订单"；')
add_body_para('• 不能改签到同一场次→"不能改签到同一场次"。')
add_body_para('所有异常均在@Transactional保护下抛出，触发事务回滚。定时任务中的异常被独立catch(log.error)，不影响其他过期订单处理。其他未知异常由GlobalExceptionHandler统一捕获并返回ApiResponse.error()。')

add_heading_styled('6.1.5 测试要点', level=3)
add_body_para('• 并发下单测试：模拟多用户同时选择相同座位，验证乐观锁正确工作（一个成功一个回滚）和事务完整性；')
add_body_para('• 支付状态测试：验证非待支付订单无法支付、他人订单无法支付、已开场场次无法支付；')
add_body_para('• 改签差价测试：验证需补差场景（新票价>旧票价→新订单待支付）和无需补差场景（新票价≤旧票价→新订单直接生效）；')
add_body_para('• 超时取消测试：验证订单在创建15分钟后被自动取消，座位和sold_count正确恢复；')
add_body_para('• 边界测试：验证已开场/已结束场次的订单操作全部被拒绝；')
add_body_para('• 事务回滚测试：模拟中间步骤失败，验证前置数据库操作全部回滚。')

# 6.2 AuthServiceImpl
add_heading_styled('6.2 软件单元2（SU_BL_01: AuthServiceImpl——认证服务实现）', level=2)

add_heading_styled('6.2.1 功能描述', level=3)
add_body_para('AuthServiceImpl负责系统的用户注册、登录认证和密码修改功能。登录时根据loginType区分用户端（查user表）和管理端（查employee表），认证成功后生成JWT Token返回。注册时进行用户名唯一性校验（同时检查user表和employee表防止跨表重名）。')

add_heading_styled('6.2.2 处理流程', level=3)
add_body_para('（1）login(LoginRequest)：根据loginType判断USER或ADMIN→查对应表（userMapper.findByUsername或employeeMapper.findByUsername）→若账号不存在或密码错误均返回"用户名或密码错误"（防用户枚举）→BCrypt.matches()密码验证→JwtTokenProvider.generateToken()生成JWT（载荷含userId/username/role/loginType，HMAC-SHA256签名，24小时有效期）→构建LoginResponse{token, userId, username, realName, roleCode, roleName, permissions, theme}。')
add_body_para('（2）register(RegisterRequest)：检查username在user表和employee表中均唯一→BCrypt.encode()加密密码→设置默认status=0, theme="white"→nickname为空则用username→插入user表。@Transactional保证原子性。')
add_body_para('（3）changePassword(Long userId, String oldPassword, String newPassword)：验证旧密码正确→BCrypt编码新密码→更新password字段。')

add_heading_styled('6.2.3 内部数据', level=3)
add_body_para('依赖UserMapper、EmployeeMapper、RoleMapper、PasswordEncoder、JwtTokenProvider。无自身状态属性。')

add_heading_styled('6.2.4 异常与错误处理', level=3)
add_body_para('• 用户名不存在→"用户名或密码错误"（与密码错误统一提示，防止攻击者枚举有效用户名）；')
add_body_para('• 密码错误→"用户名或密码错误"；')
add_body_para('• 用户名已存在→"用户名已存在"（注册时分别检查user表和employee表）；')
add_body_para('• 账号被禁用→"账号已被禁用，请联系管理员"；')
add_body_para('• 旧密码错误→"原密码不正确"。')

add_heading_styled('6.2.5 测试要点', level=3)
add_body_para('• 验证USER和ADMIN两种loginType正确路由到不同的表（user/employee）；')
add_body_para('• 验证被禁用账号(status=1)无法登录；')
add_body_para('• 验证注册时用户名重复检查覆盖user和employee两张表；')
add_body_para('• 验证修改密码后旧Token仍然有效（JWT无状态特性——由JWT过期时间控制）。')

# 6.3 ScheduleServiceImpl
add_heading_styled('6.3 软件单元3（SU_BL_03: ScheduleServiceImpl——场次服务实现）', level=2)

add_heading_styled('6.3.1 功能描述', level=3)
add_body_para('ScheduleServiceImpl负责场次的CRUD管理、座位矩阵的生成和查询、场次信息填充（批量查询优化N+1问题），以及排片时的时间冲突检测。')

add_heading_styled('6.3.2 处理流程', level=3)
add_body_para('（1）add(Schedule)：验证关联影片和影厅存在→时间冲突检测（查询同影厅+同时间段+status=1且未结束的场次→冲突判定公式：新start<旧end AND 新end>旧start→若检测到冲突则抛异常）→插入场次记录。')
add_body_para('（2）getSeats(Long scheduleId)：查场次信息→查seat表（seatMapper.selectByScheduleId）→若座位表为空则调用generateSeats()自动生成→将一维座位列表转为二维矩阵返回。generateSeats()根据Hall的rowCount×colCount创建所有座位，行号用字母表示（A-Z, AA-AB...），解析hall.seat_layout（JSON格式不可用位置如["1-5","2-10"]）标记对应座位status=3。含二次检查机制防止并发重复生成。')
add_body_para('（3）fillScheduleInfo(List<Schedule>)批量填充：一次性收集所有schedule的movieId和hallId→调用movieMapper.selectBatchIds()和hallMapper.selectBatchIds()批量查询（而非逐条查）→在内存中Map组装→计算availableSeats=总座位-已售-已锁定。此方法将N+1次查询优化为3次查询。')

add_heading_styled('6.3.3 内部数据', level=3)
add_body_para('依赖ScheduleMapper、MovieMapper、HallMapper、SeatMapper。无状态属性。')

add_heading_styled('6.3.4 异常与错误处理', level=3)
add_body_para('• 影片或影厅不存在→BusinessException；')
add_body_para('• 时间冲突→"该影厅该时段已有其他场次安排"（冲突检测返回true时）；')
add_body_para('• 座位并发重复生成→二次检查机制（if(seats.isEmpty())后再查一次）防止。')

add_heading_styled('6.3.5 测试要点', level=3)
add_body_para('• 验证时间冲突检测的边界条件（相邻不重叠场次应通过，部分重叠应拒绝）；')
add_body_para('• 验证座位延迟生成：首次访问生成，再次访问直接返回；')
add_body_para('• 验证fillScheduleInfo批量查询替代N+1次查询的性能提升。')

# 6.4 JWT组件
add_heading_styled('6.4 软件单元4（安全组件：JwtAuthenticationFilter & JwtTokenProvider）', level=2)

add_heading_styled('6.4.1 功能描述', level=3)
add_body_para('JwtAuthenticationFilter继承OncePerRequestFilter（保证每个请求只执行一次），在每次HTTP请求时从Authorization Header中提取Bearer Token，调用JwtTokenProvider解析验证后，将用户身份信息（userId作为Principal，role作为GrantedAuthority）设置到Spring Security的SecurityContext中。白名单路径（登录/注册/公开场次查询/静态资源）跳过JWT校验。')
add_body_para('JwtTokenProvider负责JWT Token的生成、解析和验证。使用HMAC-SHA256算法签名（密钥来自application.yml的jwt.secret），Token载荷包含userId、username、role、loginType，有效期24小时（jwt.expiration=86400000ms）。')

add_heading_styled('6.4.2 处理流程', level=3)
add_body_para('请求到达→shouldSkip()检查请求路径是否匹配白名单（/api/auth/login, /api/auth/register, /api/schedules/query/**, /uploads/**，使用AntPathMatcher进行路径模式匹配）→在白名单则直接放行chain.doFilter()→不在白名单则从Header提取token（去掉"Bearer "前缀）→JwtTokenProvider.validateToken()验证签名和过期时间→验证失败则直接放行（不设置认证信息，由后续SecurityConfig的URL权限规则决定是否拒绝）→验证成功则提取userId和role→构建UsernamePasswordAuthenticationToken(userId, token, [SimpleGrantedAuthority(role)])→setAuthentication到SecurityContextHolder。')

add_heading_styled('6.4.3 内部数据', level=3)
add_body_para('静态白名单路径列表：/api/auth/login, /api/auth/register, /api/schedules/query/**, /uploads/**。JWT密钥来自@Value("${jwt.secret}")。AntPathMatcher用于路径模式匹配。')

add_heading_styled('6.4.4 异常与错误处理', level=3)
add_body_para('• Token过期→ExpiredJwtException（validateToken返回false，不阻塞请求，不设置认证信息）；')
add_body_para('• Token签名无效→SignatureException（同上）；')
add_body_para('• Token格式错误→MalformedJwtException（同上）；')
add_body_para('• 无Authorization Header→直接放行chain.doFilter()，不设置认证信息（公开接口正常访问，保护接口由SecurityConfig拒绝）。')

add_heading_styled('6.4.5 测试要点', level=3)
add_body_para('• 验证白名单路径不校验Token（无需Authorization Header即可访问）；')
add_body_para('• 验证过期Token被拒绝（SecurityContext中无认证信息）；')
add_body_para('• 验证被篡改的Token（修改Payload后签名不匹配）被拒绝；')
add_body_para('• 验证认证成功后SecurityContext中userId和role正确设置。')

# 6.5 StatisticsServiceImpl
add_heading_styled('6.5 软件单元5（SU_BL_06: StatisticsServiceImpl——统计服务实现）', level=2)

add_heading_styled('6.5.1 功能描述', level=3)
add_body_para('StatisticsServiceImpl负责系统的数据统计功能，包括营收概览（总营收、总订单数、总售票数、平均票价）、每日营收趋势、近12个月月度数据、影片票房排行，以及Excel报表导出。统计数据来源于order表（过滤status IN(1,2)即待观影和已完成的订单）。')

add_heading_styled('6.5.2 处理流程', level=3)
add_body_para('• getRevenue(startDate, endDate)：按payTime范围过滤status IN(1,2)的订单→selectList加载→Java层Stream聚合计算总营收(BigDecimal求和)、总订单数、总售票数(seat_count求和)、平均票价。')
add_body_para('• getDailyRevenue(startDate, endDate)：按日分组统计每日营收，返回每日数据列表用于图表展示。')
add_body_para('• getMonthlyData()：统计最近12个月的月度营收、订单数、售票数。')
add_body_para('• getMovieRanking(limit)：按movie_id分组统计每个影片的总票房，按票房降序排列取前N名。')
add_body_para('• export(startDate, endDate, HttpServletResponse)：使用Apache POI创建XSSFWorkbook(.xlsx)→构建Sheet包含营收汇总行和订单明细行→通过response.getOutputStream()输出，设置Content-Disposition附件下载。')

add_heading_styled('6.5.3 内部数据', level=3)
add_body_para('依赖OrderMapper。统计数据使用Java Stream API和Map进行分组聚合。Excel导出使用XSSFWorkbook。')

add_heading_styled('6.5.4 异常与错误处理', level=3)
add_body_para('• 无数据时返回空统计结果（data中各字段为0）；')
add_body_para('• 日期格式错误→由GlobalExceptionHandler统一处理。')

add_heading_styled('6.5.5 测试要点', level=3)
add_body_para('• 验证无订单时返回0值而非null或异常；')
add_body_para('• 验证日期范围过滤正确（含起始日、不含结束日次日）；')
add_body_para('• 验证月度统计的近12个月范围正确；')
add_body_para('• 验证Excel导出文件格式和内容完整。')

doc.add_page_break()

# ============================================================
# 7. 开发架构设计
# ============================================================
add_heading_styled('7. 开发架构设计', level=1)

# 7.1 工程结构
add_heading_styled('7.1 工程结构', level=2)
add_body_para('本项目工程的目录结构如下：')

add_body_para('''
TTMS/
├── backend/                              # 后端Spring Boot工程
│   ├── pom.xml                           # Maven依赖配置（Spring Boot 3.2.5 + MyBatis-Plus 3.5.6等）
│   └── src/main/
│       ├── java/com/ttms/
│       │   ├── TTMSApplication.java      # 启动类
│       │   │   (@SpringBootApplication + @EnableScheduling + @EnableTransactionManagement)
│       │   ├── config/                   # 配置包（5个类）
│       │   │   ├── CorsConfig.java           # 跨域配置（允许localhost:*）
│       │   │   ├── DatabaseInitializer.java  # 数据库初始化（CommandLineRunner，启动时执行）
│       │   │   ├── MyBatisPlusConfig.java    # 分页插件（最大500条）+ 自动填充createTime/updateTime
│       │   │   ├── SecurityConfig.java       # Spring Security URL规则 + BCryptPasswordEncoder
│       │   │   └── WebMvcConfig.java         # 静态资源映射 + SPA路由回退(index.html)
│       │   ├── controller/               # 控制器包（10个类）
│       │   │   ├── AuthController.java       # 认证：登录/注册/修改密码
│       │   │   ├── MovieController.java      # 影片：公开查询 + 管理CRUD
│       │   │   ├── ScheduleController.java   # 场次：公开查询 + 管理CRUD
│       │   │   ├── OrderController.java      # 用户端订单：创建/支付/改签/退票
│       │   │   ├── AdminOrderController.java # 管理端订单：列表/协助下单
│       │   │   ├── HallController.java       # 影厅CRUD
│       │   │   ├── EmployeeController.java   # 员工CRUD + @PreAuthorize方法级权限
│       │   │   ├── StatisticsController.java # 统计：营收/排行/Excel导出
│       │   │   ├── SystemController.java     # 系统配置/操作日志/用户主题
│       │   │   └── FileController.java       # 图片上传（魔数校验）
│       │   ├── service/                  # 服务接口包（6个接口+6个实现）
│       │   │   ├── AuthService.java / impl/AuthServiceImpl.java
│       │   │   ├── MovieService.java / impl/MovieServiceImpl.java
│       │   │   ├── ScheduleService.java / impl/ScheduleServiceImpl.java
│       │   │   ├── HallService.java / impl/HallServiceImpl.java
│       │   │   ├── OrderService.java / impl/OrderServiceImpl.java
│       │   │   └── StatisticsService.java / impl/StatisticsServiceImpl.java
│       │   ├── mapper/                   # 数据访问接口包（10个Mapper）
│       │   │   ├── UserMapper.java       # 继承BaseMapper<User>
│       │   │   ├── EmployeeMapper.java   # 继承BaseMapper<Employee>
│       │   │   ├── MovieMapper.java      # +searchMovies自定义方法
│       │   │   ├── HallMapper.java
│       │   │   ├── ScheduleMapper.java   # +selectByMovieId等5个自定义方法
│       │   │   ├── SeatMapper.java       # +lockSeat/markSold/releaseSeatsByOrderId
│       │   │   ├── OrderMapper.java      # +selectByOrderNo/selectExpiredOrders等
│       │   │   ├── OrderLogMapper.java
│       │   │   ├── RoleMapper.java
│       │   │   └── SystemConfigMapper.java
│       │   ├── entity/                   # 实体类包（10个实体，均@Data + @TableName）
│       │   │   ├── User.java, Employee.java, Movie.java, Hall.java
│       │   │   ├── Schedule.java, Seat.java, Order.java, OrderLog.java
│       │   │   ├── Role.java, SystemConfig.java
│       │   ├── dto/                      # 数据传输对象包（6个DTO）
│       │   │   ├── ApiResponse.java      # 统一响应{code, message, data}+静态工厂
│       │   │   ├── LoginRequest.java     # {username, password, loginType}
│       │   │   ├── LoginResponse.java    # {token, userId, username, realName, roleCode, ...}
│       │   │   ├── RegisterRequest.java  # {username, password, phone, email, nickname}
│       │   │   ├── OrderRequest.java     # {scheduleId, seatNumbers[], userId?}
│       │   │   └── RescheduleRequest.java# {orderId, newScheduleId, newSeatNumbers[]}
│       │   ├── security/                 # 安全包（2个类）
│       │   │   ├── JwtTokenProvider.java        # JWT生成/解析/验证（jjwt 0.12.5）
│       │   │   └── JwtAuthenticationFilter.java  # OncePerRequestFilter
│       │   └── exception/                # 异常包（2个类）
│       │       ├── BusinessException.java       # RuntimeException + code字段
│       │       └── GlobalExceptionHandler.java  # @RestControllerAdvice处理10+种异常
│       └── resources/
│           ├── application.yml           # 应用配置（端口/数据库/JWT/上传/MyBatis-Plus）
│           ├── schema.sql                # 建表脚本（10张表+ALTER兼容+示例数据）
│           └── static/                   # Vue前端编译产物（SPA静态资源，约15个chunk文件）
└── frontend/                             # Vue 3前端工程源码（开发环境）
    ├── package.json                      # Node依赖
    ├── vite.config.js                    # Vite构建配置
    └── src/
        ├── App.vue                       # 根组件
        ├── main.js                       # 入口文件
        ├── router/                       # Vue Router路由配置（History模式）
        ├── stores/                       # Pinia状态管理（User Store等）
        ├── views/                        # 页面组件（约15个Vue文件）
        │   ├── Login.vue, Register.vue, Home.vue
        │   ├── MovieDetail.vue, SeatSelection.vue
        │   ├── MyOrders.vue, Profile.vue
        │   ├── Dashboard.vue, MovieManage.vue, HallManage.vue
        │   ├── ScheduleManage.vue, OrderManage.vue
        │   ├── EmployeeManage.vue, Statistics.vue, SystemSettings.vue
        ├── components/                   # 公共组件（NavBar等）
        └── api/                          # Axios API模块封装（movie.js, order.js等）
''', first_line_indent=False)

# 7.2 源代码文件
add_heading_styled('7.2 源代码文件', level=2)
add_body_para('本项目工程的源代码文件构成及其与软件逻辑单元的关系如下：')
add_body_para('表18 源代码文件与逻辑单元对应关系', bold=True, first_line_indent=False)
src_rows = [
    ['controller/AuthController.java', 'SU_UI_01', '用户认证接口'],
    ['controller/MovieController.java', 'SU_UI_02', '影片管理接口'],
    ['controller/ScheduleController.java', 'SU_UI_03', '场次管理接口'],
    ['controller/OrderController.java', 'SU_UI_04', '用户端订单接口'],
    ['controller/AdminOrderController.java', 'SU_UI_05', '管理端订单接口'],
    ['controller/HallController.java', 'SU_UI_06', '影厅管理接口'],
    ['controller/EmployeeController.java', 'SU_UI_07', '员工管理接口'],
    ['controller/StatisticsController.java', 'SU_UI_08', '数据统计接口'],
    ['controller/SystemController.java', 'SU_UI_09', '系统配置接口'],
    ['controller/FileController.java', 'SU_UI_10', '文件上传接口'],
    ['service/impl/OrderServiceImpl.java', 'SU_BL_05', '订单核心业务（约550行，最复杂）'],
    ['service/impl/ScheduleServiceImpl.java', 'SU_BL_03', '场次核心业务（约520行）'],
    ['service/impl/AuthServiceImpl.java', 'SU_BL_01', '认证核心业务（约150行）'],
    ['service/impl/StatisticsServiceImpl.java', 'SU_BL_06', '统计核心业务（约350行）'],
    ['service/impl/MovieServiceImpl.java', 'SU_BL_02', '影片核心业务'],
    ['service/impl/HallServiceImpl.java', 'SU_BL_04', '影厅核心业务'],
    ['mapper/SeatMapper.java', 'SU_DA_06', '座位数据访问（含乐观锁SQL）'],
    ['mapper/ScheduleMapper.java', 'SU_DA_05', '场次数据访问（含冲突检测SQL）'],
    ['mapper/OrderMapper.java', 'SU_DA_07', '订单数据访问（含原子操作SQL）'],
    ['security/JwtAuthenticationFilter.java', '安全组件', 'JWT请求过滤器（约130行）'],
    ['security/JwtTokenProvider.java', '安全组件', 'JWT令牌生成与验证（约110行）'],
    ['config/SecurityConfig.java', '安全组件', 'Spring Security配置'],
    ['config/DatabaseInitializer.java', '初始化组件', '数据库初始数据加载（约260行）'],
    ['exception/GlobalExceptionHandler.java', '异常处理', '全局异常统一处理（约120行）'],
]
add_table(['文件路径（相对backend/src/main/java/com/ttms/）', '对应逻辑单元', '说明'], src_rows, [7, 3, 6])

# 7.3 系统组件
add_heading_styled('7.3 系统组件', level=2)
add_body_para('本系统的组件构成及关系：')
add_body_para('（1）Spring Boot核心容器：管理所有Bean的生命周期和依赖注入。TTMSApplication为入口（@SpringBootApplication + @EnableScheduling + @EnableTransactionManagement）。')
add_body_para('（2）Spring Security过滤器链：CorsConfig（跨域处理）→JwtAuthenticationFilter（JWT认证）→SecurityConfig（URL权限规则），在请求到达Controller之前执行。')
add_body_para('（3）MyBatis-Plus组件：PaginationInnerInterceptor（分页插件，最大500条）、MetaObjectHandler（自动填充createTime/updateTime）、@TableLogic（逻辑删除自动过滤），在数据访问层透明生效。')
add_body_para('（4）Druid连接池组件：管理数据库连接（初始5个、最小5个、最大20个），提供SQL监控功能。')
add_body_para('（5）Jackson序列化组件：时区Asia/Shanghai，日期格式yyyy-MM-dd HH:mm:ss，null值不序列化（NON_NULL），@JsonProperty(access=WRITE_ONLY)防止密码泄露。')
add_body_para('（6）Apache POI组件：用于Excel报表导出（.xlsx格式）。')
add_body_para('（7）Vue前端构建产物：编译后的JS/CSS/HTML文件位于static/目录，由Spring Boot作为静态资源提供服务。WebMvcConfig确保SPA路由回退。')
add_body_para('组件协作关系：HTTP请求→Druid连接池（数据库连接）→Spring Security过滤器链（认证授权）→Spring Boot容器（依赖注入+事务管理）→Controller→Service→MyBatis-Plus→Mapper→MySQL数据库。横切关注点（日志、异常处理、定时任务）通过Spring AOP在容器层面统一管理。')

doc.add_page_break()

# ============================================================
# 8. 物理架构设计
# ============================================================
add_heading_styled('8. 物理架构设计', level=1)

# 8.1 网络环境
add_heading_styled('8.1 网络环境', level=2)
add_body_para('本系统的网络拓扑架构如图5所示（网络拓扑图——待补充）：', bold=True, first_line_indent=False)

add_body_para('''
                    ┌─────────────────────────────────┐
                    │         客户端浏览器               │
                    │  (Chrome / Firefox / Edge /       │
                    │   Safari，支持ES2020)              │
                    └──────────────┬──────────────────┘
                                   │ HTTP/HTTPS (默认端口8080)
                                   │
                    ┌──────────────▼──────────────────┐
                    │         应用服务器                  │
                    │  Spring Boot JAR (内嵌Tomcat)      │
                    │  - 静态资源服务 (Vue SPA)          │
                    │  - REST API 服务 (JSON)            │
                    │  - JWT 认证过滤                     │
                    │  - 文件上传存储 (./uploads/)        │
                    │  建议JVM: -Xms512m -Xmx1024m       │
                    └──────────────┬──────────────────┘
                                   │ JDBC (默认端口3306)
                                   │
                    ┌──────────────▼──────────────────┐
                    │      MySQL 数据库服务器             │
                    │  MySQL 8.0+                       │
                    │  数据库名: TTMS                    │
                    │  字符集: utf8mb4                   │
                    │  存储引擎: InnoDB                  │
                    │  连接池: Druid (最大20连接)         │
                    └─────────────────────────────────┘
''', first_line_indent=False)

add_body_para('相关资源和配置说明：')
add_body_para('• 应用服务器与数据库服务器可部署在同一物理机（开发/小型影院场景），也可分离部署（生产/大型影院场景）；')
add_body_para('• 客户端通过HTTP协议访问应用服务器，默认端口8080；')
add_body_para('• 应用服务器通过JDBC协议连接MySQL数据库，默认端口3306；')
add_body_para('• 生产环境建议在前端加Nginx反向代理，配置HTTPS证书和负载均衡；')
add_body_para('• 文件上传存储在应用服务器的本地磁盘./uploads/目录，生产环境可改用云存储（如阿里云OSS/腾讯云COS）以提高可靠性和扩展性。')

# 8.2 部署方案
add_heading_styled('8.2 部署方案', level=2)
add_body_para('本系统采用前后端合并部署方案（JAR包内含前端静态资源）。')
add_body_para('部署图如图6所示（部署图——待补充）：', bold=True, first_line_indent=False)
add_body_para('具体部署步骤如下：')
add_body_para('第1步——环境准备：')
add_body_para('• 安装JDK 17及以上版本（java -version确认）；')
add_body_para('• 安装MySQL 8.0及以上版本，确保MySQL服务已启动；')
add_body_para('• 数据库TTMS可手动创建，也可由应用自动创建（JDBC URL含createDatabaseIfNotExist=true参数）。')
add_body_para('第2步——构建部署包：')
add_body_para('• 执行Maven构建：cd backend && mvn clean package -DskipTests；')
add_body_para('• 构建产物：target/TTMS-0.0.1-SNAPSHOT.jar（含内嵌Tomcat + 全部依赖 + 前端静态资源，约60-80MB）。')
add_body_para('第3步——配置外部化（可选）：')
add_body_para('• 将application.yml放置在JAR包同目录，Spring Boot会自动使用外部配置覆盖JAR内默认值；')
add_body_para('• 至少应修改：数据库密码（spring.datasource.druid.password）、JWT密钥（jwt.secret）为安全的随机字符串；')
add_body_para('• 生产环境调整CORS允许的域名（CorsConfig.java需重新编译或通过配置项支持）。')
add_body_para('第4步——启动应用：')
add_body_para('• 执行：java -jar -Xms512m -Xmx1024m TTMS-0.0.1-SNAPSHOT.jar；')
add_body_para('• 启动过程中自动执行：schema.sql建表（IF NOT EXISTS确保安全）→DatabaseInitializer.run()初始化角色/管理员/系统配置/示例影片；')
add_body_para('• 控制台输出"TTMS启动成功"即表示系统已就绪。')
add_body_para('第5步——访问验证：')
add_body_para('• 顾客端首页：http://服务器IP:8080/ → 应显示影片列表页面；')
add_body_para('• 管理端：使用默认超级管理员登录（用户名admin，密码admin123）→ 登录后自动跳转管理仪表盘；')
add_body_para('• 首次登录后建议立即修改默认管理员密码。')
add_body_para('部署注意事项：')
add_body_para('• 首次启动前确保MySQL服务已运行且连接信息正确（用户名/密码/端口）；')
add_body_para('• uploads目录需要应用进程有读写权限；')
add_body_para('• 生产环境建议配置日志文件输出（logback-spring.xml），设置滚动策略防止日志文件过大；')
add_body_para('• 应用默认端口8080如被占用，可通过server.port配置项修改；')
add_body_para('• 如需升级，停旧JAR→备份数据库→部署新JAR→启动即可（schema.sql使用IF NOT EXISTS安全，DatabaseInitializer检查已有数据跳过）。')

# ============================================================
# 保存
# ============================================================
output_path = r'c:\Users\Sprite\Desktop\TTMS软件设计说明书（SDS）.docx'
doc.save(output_path)
print(f'文档已成功生成：{output_path}')
print(f'文档包含完整的8个章节、18张表格、6张图占位、20行文档修改记录。')
print(f'个人信息（姓名/学号/项目编号等）均已留空，请自行填写。')
print(f'标注"待补充"的图表位置可插入实际架构图/ER图/流程图/部署图。')
